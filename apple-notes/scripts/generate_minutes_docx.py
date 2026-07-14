#!/usr/bin/env python3
"""Generate a Chinese meeting-minutes DOCX from structured JSON."""

from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "reference.docx"
TITLE_STYLE = "Heading 1"
SECTION_STYLE = "Heading 2"
SUBSECTION_STYLE = "Heading 3"
BODY_STYLE = "Normal (Web)"


def text_run(paragraph):
    return next((run for run in paragraph.runs if run.text.strip()), None)


def capture_run_format(run) -> dict[str, Any]:
    return {
        "font_name": run.font.name,
        "size": run.font.size,
        "bold": run.bold,
    }


def capture_paragraph_profile(document: Document, style_name: str) -> dict[str, Any]:
    for paragraph in document.paragraphs:
        if paragraph.style.name == style_name and paragraph.text.strip():
            run = text_run(paragraph)
            return {
                "style": style_name,
                "alignment": paragraph.alignment,
                "run": capture_run_format(run) if run else {},
            }
    return {"style": style_name, "alignment": None, "run": {}}


def capture_profiles(document: Document) -> dict[str, dict[str, Any]]:
    return {
        TITLE_STYLE: capture_paragraph_profile(document, TITLE_STYLE),
        SECTION_STYLE: capture_paragraph_profile(document, SECTION_STYLE),
        SUBSECTION_STYLE: capture_paragraph_profile(document, SUBSECTION_STYLE),
        BODY_STYLE: capture_paragraph_profile(document, BODY_STYLE),
    }


def apply_run_format(run, profile: dict[str, Any]) -> None:
    run_format = profile.get("run") or {}
    font_name = run_format.get("font_name")
    if font_name:
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), font_name)
    if run_format.get("size") is not None:
        run.font.size = run_format["size"]
    if run_format.get("bold") is not None:
        run.bold = run_format["bold"]


def set_para(paragraph, text: str, profile: dict[str, Any]) -> None:
    style = profile.get("style")
    if style:
        paragraph.style = style
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        apply_run_format(run, profile)
    paragraph.alignment = profile.get("alignment")


def clear_body(document: Document) -> None:
    body = document._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_heading(document: Document, text: str, level: int, profiles: dict[str, dict[str, Any]]) -> None:
    style = SECTION_STYLE if level == 1 else SUBSECTION_STYLE
    p = document.add_paragraph(style=style)
    set_para(p, text, profiles[style])


def add_body_item(document: Document, text: str, profiles: dict[str, dict[str, Any]]) -> None:
    p = document.add_paragraph(style=BODY_STYLE)
    set_para(p, text, profiles[BODY_STYLE])


def normalize_filename_date(value: str | None) -> str:
    if not value:
        return date.today().isoformat()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", value):
        raise SystemExit("--date must use YYYY-MM-DD")
    return value


def safe_filename(value: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "-", value)


def build_output_prefix(minutes: dict[str, Any]) -> str:
    metadata = minutes.get("metadata") or {}
    topic = metadata.get("会议主题") or "会议"
    units_raw = metadata.get("参会单位") or ""
    units = "&".join(u.strip() for u in re.split(r"[，,、]", units_raw) if u.strip()) or "参会方"
    return f"{topic}_{units}_会议纪要"


def compact_date(value: str) -> str:
    return value.replace("-", "")


def split_units(units_raw: str) -> list[str]:
    return [u.strip() for u in re.split(r"[，,、]", units_raw) if u.strip()]


def title_units(metadata: dict[str, Any]) -> str:
    units = split_units(str(metadata.get("参会单位") or ""))
    if len(units) >= 2:
        return f"{units[0]}与{units[1]}"
    if len(units) == 1:
        return units[0]
    return "参会单位"


def title_topic(minutes: dict[str, Any]) -> str:
    metadata = minutes.get("metadata") or {}
    topic = metadata.get("会议主题")
    if topic:
        return str(topic)
    title = str(minutes.get("title") or "会议")
    return re.sub(r"会议纪要$", "", title) or "会议"


def add_metadata(document: Document, minutes: dict[str, Any], profiles: dict[str, dict[str, Any]]) -> None:
    metadata = minutes.get("metadata") or {}
    for line in [title_units(metadata), title_topic(minutes), "会议纪要"]:
        p = document.add_paragraph(style=TITLE_STYLE)
        set_para(p, line, profiles[TITLE_STYLE])

    add_heading(document, "一、会议背景", 1, profiles)
    add_body_item(document, f"会议时间：{metadata.get('会议时间', '【待补充】')}", profiles)
    add_body_item(document, f"会议地点：{metadata.get('会议地点', '【待补充】')}", profiles)
    add_body_item(document, "主要参会与发言人员", profiles)
    participants = [str(person) for person in minutes.get("participants") or ["【待补充】"]]
    for person in participants:
        add_body_item(document, person, profiles)
    host = str(metadata.get("主持人") or "").strip()
    if host and host not in {"【待补充】", "待补充", "待确认"}:
        add_body_item(document, f"主持人：{host}", profiles)


def add_sections(document: Document, sections: list[dict[str, Any]], profiles: dict[str, dict[str, Any]]) -> None:
    for section in sections:
        heading = section.get("heading")
        if heading:
            add_heading(document, str(heading), 1, profiles)

        for item in section.get("items") or []:
            if isinstance(item, dict):
                topic = item.get("topic")
                discussion = item.get("discussion")
                conclusion = item.get("conclusion")
                if topic:
                    p = document.add_paragraph(style=BODY_STYLE)
                    set_para(p, str(topic), profiles[BODY_STYLE])
                    for run in p.runs:
                        run.bold = True
                if discussion:
                    add_body_item(document, f"讨论：{discussion}", profiles)
                if conclusion:
                    add_body_item(document, f"结论：{conclusion}", profiles)
            else:
                add_body_item(document, str(item), profiles)

        for subsection in section.get("subsections") or []:
            sub_heading = subsection.get("heading")
            if sub_heading:
                add_heading(document, str(sub_heading), 2, profiles)
            for item in subsection.get("items") or []:
                add_body_item(document, str(item), profiles)


def reference_table_profile(document: Document) -> dict[str, Any]:
    if not document.tables:
        return {
            "headers": ["序号", "行动项", "具体内容", "负责人", "时间节点"],
            "widths": [704, 2340, 5036, 1185, 1500],
            "style": "Table Grid",
            "header": {"font_name": "仿宋", "bold": True},
            "body": {"font_name": "仿宋"},
        }
    table = document.tables[0]
    headers = [cell.text for cell in table.rows[0].cells]
    widths = []
    for cell in table.rows[0].cells:
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        widths.append(int(tc_w.get(qn("w:w"))) if tc_w is not None and tc_w.get(qn("w:w")) else 0)
    header_run = text_run(table.rows[0].cells[0].paragraphs[0])
    body_run = text_run(table.rows[1].cells[0].paragraphs[0]) if len(table.rows) > 1 else None
    return {
        "headers": headers,
        "widths": widths,
        "style": table.style.name,
        "header": capture_run_format(header_run) if header_run else {},
        "body": capture_run_format(body_run) if body_run else {},
    }


def apply_cell_run_format(run, run_format: dict[str, Any]) -> None:
    font_name = run_format.get("font_name")
    if font_name:
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), font_name)
    if run_format.get("size") is not None:
        run.font.size = run_format["size"]
    if run_format.get("bold") is not None:
        run.bold = run_format["bold"]


def set_table_geometry(table, table_profile: dict[str, Any]) -> None:
    widths = table_profile["widths"]
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[col_index]))

            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side in ["top", "left", "bottom", "right"]:
                margin = margins.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    margins.append(margin)
                margin.set(qn("w:w"), "60")
                margin.set(qn("w:type"), "dxa")

            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    apply_cell_run_format(run, table_profile["header" if row_index == 0 else "body"])


def action_value(action: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = action.get(key)
        if value:
            return str(value)
    return ""


def action_details(action: dict[str, Any]) -> str:
    explicit = action_value(action, "具体内容")
    if explicit:
        return explicit
    parts = []
    background = action_value(action, "讨论背景")
    output = action_value(action, "下一步/输出", "输出")
    status = action_value(action, "状态")
    if background:
        parts.append(background)
    if output:
        parts.append(f"下一步/输出：{output}")
    if status:
        parts.append(f"状态：{status}")
    return "；".join(parts)


def add_actions(
    document: Document,
    actions: list[dict[str, Any]],
    profiles: dict[str, dict[str, Any]],
    table_profile: dict[str, Any],
) -> None:
    add_heading(document, "七、下一步行动项", 1, profiles)
    headers = table_profile["headers"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = table_profile["style"]
    for col, header in enumerate(headers):
        table.cell(0, col).text = header
    if not actions:
        actions = [{
            "序号": "1",
            "行动项": "【待补充】",
            "具体内容": "【待补充】",
            "负责人": "【待补充】",
            "时间节点": "【待补充】",
        }]
    for index, action in enumerate(actions, start=1):
        cells = table.add_row().cells
        values = {
            "序号": action_value(action, "序号", "编号") or str(index),
            "行动项": action_value(action, "行动项", "待办事项"),
            "具体内容": action_details(action),
            "负责人": action_value(action, "负责人"),
            "时间节点": action_value(action, "时间节点", "截止时间"),
        }
        for col, header in enumerate(headers):
            cells[col].text = values.get(header, action_value(action, header))
    set_table_geometry(table, table_profile)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--json", required=True, type=Path, help="Path to structured minutes JSON")
    parser.add_argument("--template", default=DEFAULT_TEMPLATE, type=Path, help="Reference DOCX template")
    parser.add_argument("--output-dir", default=Path.cwd(), type=Path, help="Output directory")
    parser.add_argument("--date", default=None, help="Output date in YYYY-MM-DD; defaults to today")
    parser.add_argument("--output-name", default=None, help="Optional explicit output filename")
    args = parser.parse_args()

    minutes = json.loads(args.json.read_text(encoding="utf-8"))
    document = Document(args.template)
    profiles = capture_profiles(document)
    table_profile = reference_table_profile(document)
    clear_body(document)
    add_metadata(document, minutes, profiles)

    sections = deepcopy(minutes.get("sections") or [])
    actions = minutes.get("actions") or []
    content_sections = []
    conclusion_sections = []
    for section in sections:
        heading = str(section.get("heading", ""))
        if heading.startswith(("七", "五、待办")):
            continue
        if heading.startswith("八") or "会议结论" in heading:
            conclusion_sections.append(section)
        else:
            content_sections.append(section)

    add_sections(document, content_sections, profiles)
    add_actions(document, actions, profiles, table_profile)
    add_sections(document, conclusion_sections, profiles)

    out_date = normalize_filename_date(args.date)
    prefix = build_output_prefix(minutes)
    output_name = args.output_name or f"{prefix}_{compact_date(out_date)}.docx"
    output_path = args.output_dir / safe_filename(output_name)
    document.core_properties.title = minutes.get("title") or "会议纪要"
    document.save(output_path)
    print(output_path.resolve())


if __name__ == "__main__":
    main()
