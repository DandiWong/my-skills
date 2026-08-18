import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "generate_minutes_docx.py"
ZHONGSHAN = ROOT / "assets" / "template_zhongshan.docx"
INTERNAL = ROOT / "assets" / "template_internal.docx"
SKILL_MD = ROOT / "SKILL.md"
MANIFEST = ROOT / "manifest.json"
CHECK_ENV = ROOT / "scripts" / "check_env.py"


def read_frontmatter(path):
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    if not lines or lines[0] != "---":
        raise AssertionError("missing YAML frontmatter")
    frontmatter = {}
    for line in lines[1:]:
        if line == "---":
            return frontmatter
        if ":" in line:
            key, value = line.split(":", 1)
            frontmatter[key.strip()] = value.strip().strip('"')
    raise AssertionError("unterminated YAML frontmatter")


def first_run_with_text(paragraph):
    return next(run for run in paragraph.runs if run.text.strip())


def run_generator(payload, template=None):
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        json_path = tmp_path / "minutes.json"
        json_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        command = [
            sys.executable,
            str(SCRIPT),
            "--json",
            str(json_path),
            "--output-dir",
            str(tmp_path),
            "--date",
            "2026-06-11",
        ]
        if template:
            command.extend(["--template", str(template)])
        result = subprocess.run(command, check=True, capture_output=True, text=True)
        return Document(Path(result.stdout.strip()))


class GenerateMinutesDocxTests(unittest.TestCase):
    def sample_payload(self):
        return {
            "title": "项目沟通会会议纪要",
            "metadata": {
                "会议主题": "项目沟通会",
                "会议时间": "2026-06-11 10:00-11:00",
                "会议地点": "线上会议",
                "参会单位": "复星医药，合作方",
                "主持人": "张三",
            },
            "participants": ["张三（复星医药/项目负责人）", "李四（合作方/技术负责人）"],
            "sections": [
                {"heading": "二、关键结论和共识", "items": ["双方确认后续以纪要待办表跟踪事项。"]},
                {
                    "heading": "三、详细讨论要点",
                    "subsections": [
                        {"heading": "（一）项目范围", "items": ["1. 需进一步补充业务范围边界。"]}
                    ],
                },
                {
                    "heading": "四、争议项",
                    "items": [
                        {
                            "topic": "（一）时间安排",
                            "viewpoints": ["复星医药：期望两周内完成交付。", "合作方：建议三周以确保质量。"],
                            "conclusion": "截止时间待确认。",
                        }
                    ],
                },
                {
                    "heading": "六、会议总结",
                    "items": ["本次会议双方就项目推进达成共识，明确了后续分工与交付节奏。"],
                },
            ],
            "actions": [
                {
                    "编号": "1",
                    "待办事项": "补充项目范围说明",
                    "讨论背景": "会议需明确边界",
                    "下一步/输出": "范围说明文档",
                    "负责人": "张三",
                    "截止时间": "待确认",
                    "状态": "待执行",
                }
            ],
        }

    def test_generated_paragraph_styles_follow_reference_docx(self):
        reference = Document(ZHONGSHAN)
        generated = run_generator(self.sample_payload())

        self.assertEqual([p.style.name for p in generated.paragraphs[:3]], ["Heading 1"] * 3)
        self.assertEqual(generated.paragraphs[0].text, "复星医药与合作方")
        self.assertEqual(generated.paragraphs[1].text, "项目沟通会")
        self.assertEqual(generated.paragraphs[2].text, "会议纪要")

        for generated_index, reference_index in [(0, 0), (3, 3), (4, 4)]:
            generated_para = generated.paragraphs[generated_index]
            reference_para = reference.paragraphs[reference_index]
            generated_run = first_run_with_text(generated_para)
            reference_run = first_run_with_text(reference_para)
            self.assertEqual(generated_para.style.name, reference_para.style.name)
            self.assertEqual(generated_para.alignment, reference_para.alignment)
            self.assertEqual(generated_run.font.name, reference_run.font.name)
            self.assertEqual(generated_run.font.size, reference_run.font.size)

    def test_action_table_layout_uses_reference_headers_and_legacy_mapping(self):
        reference = Document(ZHONGSHAN)
        generated = run_generator(self.sample_payload())

        reference_headers = [cell.text for cell in reference.tables[0].rows[0].cells]
        generated_table = generated.tables[0]
        generated_headers = [cell.text for cell in generated_table.rows[0].cells]

        self.assertEqual(generated_headers, reference_headers)
        self.assertEqual(len(generated_table.columns), len(reference.tables[0].columns))
        self.assertEqual(
            [cell.text for cell in generated_table.rows[1].cells],
            ["1", "补充项目范围说明", "会议需明确边界；下一步/输出：范围说明文档；状态：待执行", "张三", "待确认"],
        )

    def test_skill_frontmatter_declares_semver_version(self):
        frontmatter = read_frontmatter(SKILL_MD)

        self.assertEqual(frontmatter["name"], "fosunpharma-meeting-minutes")
        self.assertEqual(frontmatter["version"], "0.4.5")
        self.assertIn("v0.4.5", frontmatter["description"])

    def test_named_and_path_templates_preserve_expected_header_logos(self):
        payload = self.sample_payload()

        self.assertEqual(len(run_generator(payload, "zhongshan").sections[0].header._element.xpath(".//a:blip")), 2)
        self.assertEqual(len(run_generator(payload, "internal").sections[0].header._element.xpath(".//a:blip")), 1)
        self.assertEqual(len(run_generator(payload, INTERNAL).sections[0].header._element.xpath(".//a:blip")), 1)

    def test_internal_template_uses_single_organization_title_and_participant_list(self):
        generated = run_generator(self.sample_payload(), "internal")
        paragraphs = [p.text for p in generated.paragraphs if p.text.strip()]

        self.assertEqual(generated.paragraphs[0].text, "复星医药")
        self.assertIn(
            "参会人员：张三（复星医药/项目负责人）；李四（合作方/技术负责人）",
            paragraphs,
        )
        self.assertNotIn("甲方", "\n".join(paragraphs))
        self.assertNotIn("乙方", "\n".join(paragraphs))

    def test_missing_template_exits_with_clear_error(self):
        with tempfile.TemporaryDirectory() as tmp:
            json_path = Path(tmp) / "minutes.json"
            json_path.write_text(json.dumps(self.sample_payload(), ensure_ascii=False), encoding="utf-8")
            result = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT),
                    "--json",
                    str(json_path),
                    "--template",
                    "missing-template",
                ],
                capture_output=True,
                text=True,
            )

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("not found", result.stderr)

    def test_participants_are_written_as_separate_lines_and_unknown_host_is_omitted(self):
        payload = self.sample_payload()
        payload["metadata"]["主持人"] = "【待补充】"
        payload["participants"] = [
            "复星医药：张三、李四",
            "中山医院：王五、赵六",
        ]

        generated = run_generator(payload)
        paragraphs = [p.text for p in generated.paragraphs if p.text.strip()]

        self.assertIn("主要参会与发言人员", paragraphs)
        self.assertIn("复星医药：张三、李四", paragraphs)
        self.assertIn("中山医院：王五、赵六", paragraphs)
        self.assertNotIn("主持人：【待补充】", paragraphs)

    def test_multi_party_title_joins_three_or_more_units(self):
        payload = self.sample_payload()
        payload["metadata"]["参会单位"] = "复星医药，中山医院，语言桥"
        payload["participants"] = [
            "复星医药：张三（产品经理）",
            "中山医院：王五（科室主任）",
            "语言桥：赵六",
        ]

        generated = run_generator(payload)
        self.assertEqual(generated.paragraphs[0].text, "复星医药、中山医院与语言桥")

    def test_manifest_matches_skill_frontmatter_and_declared_files_exist(self):
        frontmatter = read_frontmatter(SKILL_MD)
        manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))

        self.assertEqual(manifest["name"], frontmatter["name"])
        self.assertEqual(manifest["version"], frontmatter["version"])
        self.assertEqual(manifest["entry"], "SKILL.md")
        self.assertFalse(manifest["breaking"])
        for file_name in manifest["files"]:
            self.assertTrue((ROOT / file_name).exists(), file_name)

    def test_environment_check_reports_python_docx_available(self):
        result = subprocess.run(
            [sys.executable, str(CHECK_ENV)],
            cwd=ROOT,
            check=True,
            capture_output=True,
            text=True,
        )

        self.assertIn("python:", result.stdout)
        self.assertIn("python-docx: OK", result.stdout)

    def test_viewpoints_are_rendered_without_numbering(self):
        generated = run_generator(self.sample_payload())
        paragraphs = [p.text for p in generated.paragraphs if p.text.strip()]

        self.assertIn("复星医药：期望两周内完成交付。", paragraphs)
        self.assertIn("合作方：建议三周以确保质量。", paragraphs)
        self.assertNotIn("1. 复星医药：期望两周内完成交付。", paragraphs)
        self.assertNotIn("2. 合作方：建议三周以确保质量。", paragraphs)
        self.assertIn("结论：截止时间待确认。", paragraphs)

    def test_table_columns_fill_content_width(self):
        from docx.oxml.ns import qn as _qn

        generated = run_generator(self.sample_payload())
        section = generated.sections[0]
        content_width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips

        table = generated.tables[0]
        grid = table._tbl.tblGrid
        cols = grid.findall(_qn("w:gridCol"))
        widths = [int(col.get(_qn("w:w"))) for col in cols]
        self.assertEqual(sum(widths), content_width)

    def test_section_one_is_skipped_and_six_is_last(self):
        payload = self.sample_payload()
        payload["sections"] = [
            {"heading": "一、会议目的", "items": ["should not appear"]},
            {"heading": "二、关键结论和共识", "items": ["content"]},
            {"heading": "六、会议总结", "items": ["conclusion paragraph"]},
        ]
        generated = run_generator(payload)
        headings = [p.text for p in generated.paragraphs if p.style and p.style.name == "Heading 2"]

        self.assertNotIn("一、会议目的", headings)
        self.assertEqual(headings[0], "一、会议背景")
        self.assertEqual(headings[-1], "六、会议总结")
        self.assertEqual(headings[-2], "五、下一步行动项")


if __name__ == "__main__":
    unittest.main()
